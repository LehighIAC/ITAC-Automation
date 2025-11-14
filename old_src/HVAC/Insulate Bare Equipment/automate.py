"""
This script is used to generate the IAC recommendation for Install Bare Equipment
"""

import json5, sys, os, num2words
from docx import Document
from easydict import EasyDict
from python_docx_replace import docx_replace, docx_blocks
sys.path.append(os.path.join('..', '..'))
from Shared.IAC import *
from docxcompose.composer import Composer
import numpy as np
import fractions

# Load utility cost
jsonDict = json5.load(open(os.path.join('..', '..', 'Utility.json5')))
# Load database
jsonDict.update(json5.load(open('database.json5')))
# Convert to easydict
iac = EasyDict(jsonDict)

## Validate the length of all lists
N = iac.N
for i in iac:
    if isinstance(iac[i], list):
        # If the length is not N, throw an error
        if len(iac[i]) != N:
            raise Exception('Length of {0} is not {1}.'.format(i, N))
## Covert to numpy array for element-wise operations
nplist = ['TEMP', 'PTEMP', 'AMB', 'HR', 'DY', 'WK', 'SFA', 'AMB']
for i in nplist:
    iac[i] = np.array(iac[i])
## Constants
# Combined convective and radiative heat transfer coefficient; BTU/hr/Ft^2/Farenheit
h = 0.8 
# Conversion constant; Btu/hr
C1 = 0.000293 

## Calculations
# Operating Hours
iac.OH = iac.HR * iac.DY * iac.WK
# Temperature difference
iac.TD = iac.TEMP - iac.AMB
iac.PTD = iac.PTEMP - iac.AMB
# Annual Heat Loss
iac.AHL = np.rint(h * C1 * iac.SFA * (iac.TD - iac.PTD) * iac.OH)

## Savings
# Annual electricity savings
iac.ES = np.sum(iac.AHL)
# Annual demand savings
iac.DS = np.round(np.sum(iac.AHL/iac.OH),1)
# Annual cost savings
iac.ECS = round(iac.ES * iac.EC)
iac.DCS = round(iac.DS * iac.DC)
iac.ACS = iac.ECS + iac.DCS

## Implementation cost Estimate
# Labor cost
iac.LAB = np.array([iac.LABOR] * N)
iac.EST = np.add(iac.COST, iac.LAB)
# Installation cost
iac.IC = np.sum(iac.SFA * iac.EST)
# Rebate
iac.PB = payback(iac.ACS, iac.IC.item())

# Number to words
iac.AMT = num2words.num2words(N)
# Combine word to make temperature into a sentence
iac.TEMPS = [''] * N
for i in range(N):
  iac.TEMPS[i] = str(iac.TEMP[i]) + ' °F'
iac.TEMPS = combine_words(iac.TEMPS)

# Function to convert string to fraction
def convert_fraction(n):
  frac = fractions.Fraction(n)
  # Whole number
  if frac.denominator == 1:
    myFrac = frac.numerator // frac.denominator
    return str(myFrac)
  # Mixed number
  elif frac.numerator > frac.denominator:
    myFrac = frac.numerator // frac.denominator
    myFrac_str = str(myFrac) + '-' + str(frac.numerator % frac.denominator) + '/' + str(frac.denominator)
    return myFrac_str
  # Standard fraction
  myFrac_str = str(frac.numerator) + '/' + str(frac.denominator)
  return myFrac_str

## Format strings
# set to 3 digits accuracy
iac = dollar(['EC'],iac,3)
# set to 2 digits accuracy
iac = dollar(['DC'],iac,2)
# set the rest to integer
varList = ['DCS', 'ECS', 'ACS', 'IC']
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

# Create list for fractions
iac.SIZEStr = [''] * N
for i in range(N):
  iac.SIZEStr[i] = convert_fraction(iac.SIZE[i])

# Create document for each area
for i in range(N):
  iacsub = EasyDict()
  iacsub.i = str(i+1)
  # For any list or ndarray in iac, add corresponding values to iacsub
  for j in iac:
    if isinstance(iac[j], list) or isinstance(iac[j], np.ndarray):
      iacsub[j] = iac[j][i]
  # Import individual area template
  doc = Document('template 2.docx')
  # Constants that appear once
  if i == 0:
    docx_blocks(doc, single = True)
  else:
    docx_blocks(doc, single = False)
  # Replacing keys
  docx_replace(doc, **iacsub)
  # Save file as temp{i}.docx
  doc.save('tmp'+iacsub.i+'.docx')

# Import opening template
doc = Document('template 1.docx')
# Replacing keys
docx_replace(doc, **iac)
# Save file as temp0.docx
doc.save('tmp0.docx')

# Import ending template
doc = Document('template 3.docx')

# Create list for installation sentence
iac.INSTALL = []
# get the index of unique COSTs
unique, ind = np.unique(iac.COST, return_index=True)
# deduplicate COST and SIZEStr
for i in ind:
    # distinguish a/an
    vowel = num2words.num2words(iac.COST[i])
    if vowel[0] in ['a', 'e', 'i', 'o', 'u']:
        tmpstr = "an"
    else:
        tmpstr = "a"
    tmpstr += f" ${iac.COST[i]}/ft² for {iac.SIZEStr[i]} in insulation blanket"
    # captialize the first letter of the first sentence
    iac.INSTALL.append(tmpstr)
iac.INSTALL = combine_words(iac.INSTALL)

# Replacing keys
docx_replace(doc, **iac)
# Save file as temp{N+1}.docx
doc.save('tmp'+str(N+1)+'.docx')

# Combine all docx files
master = Document("tmp0.docx")
composer = Composer(master)
for i in range(N+1):
    doc_tmp = Document('tmp'+str(i+1)+'.docx')
    composer.append(doc_tmp)

savefile(composer, str(iac.REC))

# delete temp files
for i in range(N+2):
    filename = 'tmp'+str(i)+'.docx'
    os.remove(filename)
# Caveats
caveat("Please change implementation cost references if necessary.")