"""
This script is used to generate the IAC recommendation for Repair Leaks in Compressed Air Lines.
"""

import json5, sys, os
from docx import Document
from easydict import EasyDict
from python_docx_replace import docx_replace, docx_blocks
sys.path.append('..')
from Shared.IAC import *
import numpy as np
from num2words import num2words
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load config file and convert everything to EasyDict
jsonDict = json5.load(open('Repair Leaks.json5'))
jsonDict.update(json5.load(open(os.path.join('..', 'Utility.json5'))))
iac = EasyDict(jsonDict)

# Constants
PA = 14.7 # Atmosphere, psia
C1 = 28.37 # Isentropic sonic volumetric flow constant
C2 = 60.0 # Conversion constant; sec/min
C3 = 144.0 # Conversion constant; in2/ft2
C4 = 3.03e-5 # Conversion constant; HP.min/ft.lb
C5 = 0.746 # Conversion factor; kW/HP
CD = 0.8 # Coefficient of discharge for square edged orifice
k = 1.4 # Specific heat ratio of air

# Calculations
iac.OH = iac.HR * iac.DY * iac.WK
iac.RT = round(PA / iac.P0, 4)
iac.VF0 = np.pi / 4 * (iac.T0 + 460) * iac.P1 / PA * C1 * C2 * CD / C3 / np.sqrt(iac.T1 + 460)
# Number of leaks
NL = np.array([iac.NL1, iac.NL2, iac.NL3, iac.NL4, iac.NL5, iac.NL6])
# Leak diameters
LD = np.array([1.0/64, 1.0/32, 1.0/16, 1.0/8, 3.0/16, 1.0/4])
# Leak strings
LS = ["1/64", "1/32", "1/16", "1/8", "3/16", "1/4"]
# Flow rate (cfm)
FR = LD * LD * iac.VF0
# Power Loss (hp)
PL = PA * C3 * FR * k/(k-1.0) * iac.N * C4 * \
    (np.power(iac.P0/PA,(k-1.0)/(k*iac.N)) - 1.0) / (iac.EA * iac.EM)
# Demand Loss (kW/yr)
DL = PL * C5 * iac.CF * 12
# Energy Loss (kWh/yr)
EL = PL * C5 * iac.OH
# Leak Cost ($/yr)
LC = DL * iac.DC + EL * iac.EC
# Add Table 2
DS = NL * DL
ES = NL * EL
CS = NL * LC
# Convert from numpy dtype to EasyDict
iac.SNL = sum(NL).item()
iac.ADS = round(sum(DS).item())
iac.AES = round(sum(ES).item())
iac.ACS = round(sum(CS).item())

# Implementation
# Estimate 1+1 hour per leak
iac.FLC = (1+1) * iac.SNL * iac.LR
iac.IC = iac.FLC + iac.USLD
iac.PB  = payback(iac.ACS, iac.IC)

# String formatting
# eg, 'six 1/16-inch, six 1/8-inch and three 3/16-inch'
# Make a list of strings
LeakString = []
for i in range(NL.size):
    if NL[i]!=0:
        LeakString.append(num2words(NL[i]) + ' ' + LS[i] + '-inch')
iac.LeakString = combine_words(LeakString)

## Format strings
# set electricity cost to 3 digits accuracy
iac = dollar(['EC'],iac,3)
# set the natural gas and demand to 2 digits accuracy
iac = dollar(['NGC', 'DC'],iac,2)
# set the rest to integer
varList = ['LR', 'FLC', 'USLD', 'IC', 'ACS']
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

# Import docx template
doc = Document('Repair Leaks in Compressed Air Lines.docx')

# Replacing keys
docx_replace(doc, **iac)

# Add numbers to table 2
table2 = doc.tables[2]
for i in range(NL.size):
    row = table2.rows[i+1].cells
    row[0].text = LS[i]
    row[1].text = f'{round(FR[i],2):,}'
    row[2].text = f'{round(PL[i],2):,}'
    row[3].text = f'{round(DL[i],1):,}'
    row[4].text = f'{round(EL[i]):,}'
    row[5].text = f'{round(LC[i]):,}'
    # Set alignment and line spacing
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].paragraph_format.line_spacing = 1.5

# Add numbers to table 3
table3 = doc.tables[3]
for i in range(NL.size):
    row=table3.rows[i+1].cells
    row[1].text = f'{NL[i]:,}'
    row[2].text = LS[i]
    row[3].text = f'{round(DS[i],1):,}'
    row[4].text = f'{round(ES[i]):,}'
    row[5].text = f'{round(CS[i]):,}'
    # Set alignment and line spacing
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].paragraph_format.line_spacing = 1.5
# Remove rows with zero leaks
for i in reversed(range(NL.size)):
    if NL[i]==0:
        table3._tbl.remove(table3.rows[i+1]._tr)

filename = 'AR'+iac.AR+'.docx'
doc.save(os.path.join('..', 'ARs', filename))

# Caveats
caveat("Please change implementation cost references if necessary.")