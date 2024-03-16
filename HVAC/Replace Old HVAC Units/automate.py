"""
This script is used to generate the IAC recommendation for Install programmable thermostats
"""

import json5, sys, os
from docx import Document
from easydict import EasyDict
from python_docx_replace import docx_replace, docx_blocks
sys.path.append(os.path.join('..', '..'))
from Shared.IAC import *
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load utility cost
jsonDict = json5.load(open(os.path.join('..', '..', 'Utility.json5')))
# Load database
jsonDict.update(json5.load(open('database.json5')))
# Convert to easydict
iac = EasyDict(jsonDict)

## Calculations
# Maintendance Factor
if iac.FM == True:
  iac.M = 0.01
else:
  iac.M = 0.03
# Convert to numpy array
iac.TON = np.array(iac.TON)
iac.SIZE = iac.TON * 12000
iac.AGE = np.array(iac.AGE)
iac.EERB = np.array(iac.EERB)
iac.EERC = np.round(iac.EERB * (1 - iac.M) ** np.fmin(iac.AGE, 15), 1)
iac.EERP = np.array(iac.EERP)
# Total Values
iac.TTON = np.sum(iac.TON).item()
iac.CC = np.sum(iac.SIZE).item()
# Electrical Demand
iac.CED = round(np.sum((iac.SIZE/1000 * iac.LF/100)/iac.EERC), 1)
iac.PED = round(np.sum((iac.SIZE * 0.001 * iac.LF/100)/iac.EERP), 1)
# Operating hours
iac.OH = iac.HR * iac.DY * iac.WK
# Power Reduction
iac.PR = round(iac.CED - iac.PED, 1)
# Savings
iac.ES = round(iac.PR * iac.OH)
iac.DS = round(iac.PR * iac.CF/100 * iac.CS)
iac.ECS = round(iac.EC * iac.ES)
iac.DCS = round(iac.DC * iac.DS)
iac.ACS = iac.ECS + iac.DCS

iac.IC = round(iac.TTON * iac.UC)

# Rebate
iac = rebate(iac)

## Format strings
# set to 3 digits accuracy
iac = dollar(['EC', 'ERR'],iac,3)
# set to 2 digits accuracy
iac = dollar(['DC'],iac,2)
# set the rest to integer
varList = ['UC', 'ACS', 'IC', 'ECS', 'DCS', 'RB', 'MIC', 'MRB']
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

# Import docx template
doc = Document('template.docx')

# Replacing keys
docx_replace(doc, **iac)

## Adding table
table = doc.tables[1]
for i in range(len(iac.TON)):
    row=table.rows[i+1].cells
    row[0].text = iac.AREA[i]
    row[1].text = iac.TON[i]
    row[2].text = iac.SIZE[i]
    row[3].text = iac.AGE[i]
    row[4].text = iac.EERB[i]
    row[5].text = iac.EERC[i]
    row[6].text = iac.EERP[i]
    # Set alignment and line spacing
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# Remove rows with zero HVAC
for i in reversed(range(len(iac.TON), 9)):
  table._tbl.remove(table.rows[i+1]._tr)

docx_blocks(doc, mtrue = iac.FM)
docx_blocks(doc, mfalse = not iac.FM)
docx_blocks(doc, REBATE = iac.REB)

savefile(doc, iac.REC)

# Caveats
caveat("Please change implementation cost references if necessary.")