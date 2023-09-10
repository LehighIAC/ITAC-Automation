"""
This script is used to generate the IAC recommendation for Repair Leaks in Compressed Air Lines.
"""

import json5, sys, os, locale
from docx import Document
from python_docx_replace import docx_replace, docx_blocks
# Get the path of the current script
script_path = os.path.dirname(os.path.abspath(__file__))
sys.path.append(os.path.join(script_path, '..', 'shared'))
from IAC import *
import numpy as np
# Might needs to be installed
from num2words import num2words
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import docx template
doc = Document(os.path.join(script_path, 'Repair Leaks in Compressed Air Lines.docx'))
# Load config file and convert everything to local variables
iacDict = json5.load(open(os.path.join(script_path, 'Repair Leaks.json5')))
locals().update(iacDict)

# Calculations
RT = round(PA / P0, 4)
VF0 = np.pi / 4 * (T0 + 460) * P1 / PA * C1 * C2 * CD / C3 / np.sqrt(T1 + 460)
# Number of leaks
NL = np.array([NL1, NL2, NL3, NL4, NL5, NL6])
# Leak diameters
LD = np.array([1.0/64, 1.0/32, 1.0/16, 1.0/8, 3.0/16, 1.0/4])
# Leak strings
LS = ["1/64", "1/32", "1/16", "1/8", "3/16", "1/4"]
# Flow rate (cfm)
FR = LD * LD * VF0
# Power Loss (hp)
PL = PA * C3 * FR * k/(k-1.0) * N * C4 * (np.power(P0/float(PA),(k-1.0)/(k*N)) - 1.0) / (EA * EM)
# Demand Loss (kW/yr)
DL = PL * C5 * CF * 12
# Energy Loss (kWh/yr)
EL = PL * C5 * OH
# Leak Cost ($/yr)
LC = DL * DC + EL * EC
# Add Table 2
DS = NL * DL
ES = NL * EL
CS = NL * LC
# Convert from numpy dtype to python native type
SNL = sum(NL).item()
ADS = round(sum(DS).item())
AES = round(sum(ES).item())
ACS = round(sum(CS).item())

# Implementation
# Estimate 1+1 hour per leak
FLC = (1+1) * SNL * LR
IC = FLC + USLD
iacDict['PB']  = payback(ACS, IC)

# String formatting
# eg, 'six 1/16-inch, six 1/8-inch and three 3/16-inch'
LeakString = ""
count = 0
for i in range(NL.size):
    if NL[i]!=0:
        LeakString = LeakString + num2words(NL[i]) + ' ' + LS[i] + '-inch'
        count += 1
        if count <= np.count_nonzero(NL) - 2:
            LeakString = LeakString + ', '
        if count == np.count_nonzero(NL) - 1:
            LeakString = LeakString + ' and ' 
# Add leakstring to iacDict
iacDict['LeakString'] = LeakString

# Formatting
# Add all numbers in local variables to iacDict
iacDict.update({key: value for (key, value) in locals().items() if type(value) == int or type(value) == float})

# Format numbers to string with thousand separator
iacDict = grouping_num(iacDict)

# set locale to US
locale.setlocale(locale.LC_ALL, 'en_US')

# set 3 digits accuracy for electricity cost
locale._override_localeconv={'frac_digits':3}
iacDict['EC'] = locale.currency(EC, grouping=True)

# set the natural gas and demand to 2 digits accuracy
locale._override_localeconv={'frac_digits':2}
iacDict['NGC'] = locale.currency(NGC, grouping=True)
iacDict['DC'] = locale.currency(DC, grouping=True)

# set the rest to integer
locale._override_localeconv={'frac_digits':0}
for cost in ['LR', 'FLC', 'USLD', 'IC', 'ACS']:
    iacDict[cost] = locale.currency(eval(cost), grouping=True)


# Replacing keys
docx_replace(doc, **iacDict)

# Add numbers to table 2
table2 = doc.tables[2]
for i in range(NL.size):
    row = table2.rows[i+1].cells
    row[0].text = LS[i]
    row[1].text = f'{round(FR[i],2):,}'
    row[2].text = f'{round(PL[i],2):,}'
    row[3].text = f'{round(DL[i],1):,}'
    row[4].text = f'{round(EL[i]):,}'
    row[5].text = f'{round(LC[i]):,}'
    # Set alignment and line spacing
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].paragraph_format.line_spacing = 1.5

# Add numbers to table 3
table3 = doc.tables[3]
for i in range(NL.size):
    row=table3.rows[i+1].cells
    row[1].text = f'{NL[i]:,}'
    row[2].text = LS[i]
    row[3].text = f'{round(DS[i],1):,}'
    row[4].text = f'{round(ES[i]):,}'
    row[5].text = f'{round(CS[i]):,}'
    # Set alignment and line spacing
    for cell in row:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].paragraph_format.line_spacing = 1.5
# Remove rows with zero leaks
for i in reversed(range(NL.size)):
    if NL[i]==0:
        table3._tbl.remove(table3.rows[i+1]._tr)

filename = 'AR'+iacDict['AR']+'.docx'
doc.save(os.path.join(script_path, '..', 'ARs', filename))

# Caveats
print("Please change implementation cost references if necessary.")