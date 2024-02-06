"""
Fully-automated IAC report compiler
Usage: Copy all recommendations into the Recommendations folder, Update info in Compiler.json5 and Utility.json5,
then run this script.
"""


import json5, os, locale, datetime, math, platform
import pandas as pd
from easydict import EasyDict
from docx import Document, shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docxcompose.composer import Composer
from python_docx_replace import docx_replace, docx_blocks
from Shared.IAC import *

# If Recommendations/Sorted/ folder doesn't exist, create one
os.makedirs(os.path.join('Recommendations', 'Sorted'), exist_ok=True)

# If on macOS
if os.path.exists(os.path.join('Energy Charts', 'Energy Charts.fld')):
    chartPath = os.path.join('Energy Charts', 'Energy Charts.fld')
# If on Windows
elif os.path.exists(os.path.join('Energy Charts', 'Energy Charts_files')):
    chartPath = os.path.join('Energy Charts', 'Energy Charts_files')
else:
    # If chart html folder doesn't exist, exit
    raise Exception("Chart images not found. Please save the energy chart as web page (.htm).")

# Load config file and convert everything to local variables
print("Reading json5 database...", end ="")
jsonDict = json5.load(open('Compiler.json5'))
jsonDict.update(json5.load(open('Utility.json5')))
iac = EasyDict(jsonDict)
print("done")

# Initialize dataframe
columns = ["isAdditional", "File Name", "ARC No.", "Description", "Electricity (kWh)", "Electricity (MMBtu)", "Demand (kW)"
           , "Natural Gas (MMBtu)", "Other Energy Type", "Other Energy Amount", "Other Resource Type", "Other Resource Amount"
           , "Savings Type", "Savings Value", "Annual Cost Savings", "Implementation Cost", "Payback Period"]
df = pd.DataFrame(columns=columns)

# Set locale to en_US
locale.setlocale(locale.LC_ALL, 'en_US')

print("Reading recommendations...")
# Get all .docx files in Recommendations/ directory and extract information
recList = [f for f in os.listdir('Recommendations') if f.endswith('.docx')]
recID = 0
for recDoc in recList:
    print(recDoc)
    doc = Document(os.path.join('Recommendations', recDoc))
    recInfo = {}
    # Record file name
    recInfo['File Name'] = recDoc

    # Parse document title
    fullTitle = doc.paragraphs[0].text
    separatorFlag = False
    # list of possible separators
    separatorList = [":", "-", "–"]
    for separator in separatorList:
        if separator in fullTitle:
            separatorFlag = True
            # check if the document is an additional recommendation by title
            # Keep "AAR" for outdated documents
            recInfo['isAdditional'] = ("Additional" in fullTitle.split(separator)[0]) or ("AAR" in fullTitle.split(separator)[0])
            # Parse the title of the .docx file
            recInfo['Description'] = title_case(fullTitle.split(separator)[1].strip())
            break
    if separatorFlag == False:
        raise Exception("Can't parse document title:\n" + fullTitle)
    
    # Read the 1st table in .docx files
    try:
        table = doc.tables[0]
    except:
        raise Exception("Error: " + recDoc + " is not a valid recommendation. Please check if the summary table is present.")

    for row in table.rows:
        key = row.cells[0].text
        value = row.cells[1].text
        # Parse ARC Number
        if "arc" in key.lower() and "number" in key.lower():
            validate_arc(value)
            recInfo['ARC No.'] = value
        # Parse Annual Cost Savings
        elif "annual" in key.lower() and "cost" in key.lower():
            # convert currency to interger
            recInfo['Annual Cost Savings'] = locale.atoi(value.strip("$"))
        # Parse Implementation Cost
        elif "implementation" in key.lower():
            # convert currency to interger
            recInfo['Implementation Cost'] = locale.atoi(value.strip("$"))
        # If Payback Period skip (Doesn't matter, will calculate later)
        elif "payback" in key.lower():
            continue
        # Parse Electricity
        elif "electricity" in key.lower():
            recInfo['Electricity (kWh)'] = locale.atoi(value.split(' ')[0])
        # Parse Demand
        elif "demand" in key.lower():
            recInfo['Demand (kW)'] = locale.atoi(value.split(' ')[0])
        # Parse Natural Gas
        elif "natural" in key.lower():
            recInfo['Natural Gas (MMBtu)'] = locale.atoi(value.split(' ')[0])
        # Parse undefined type
        else:
            # If the value contains mmbtu, parse it as other energy
            if "mmbtu" in value.lower():
                # Remove "annual" (usually the first word)
                if "annual" in key.lower():
                    key = key.split(' ', 1)[1]
                # Remove "savings" (usually the last word)
                if "saving" in key.lower():
                    key = key.rsplit(' ', 1)[0]
                recInfo['Other Energy Type'] = title_case(key)
                # Parse number
                recInfo['Other Energy Amount'] = locale.atoi(value.split(' ')[0])
            # If not, parse it as other resource
            else:
                # Remove "annual" (usually the first word)
                if "annual" in key.lower():
                    key = key.split(' ', 1)[1]
                # Remove "savings" (usually the last word)
                if "saving" in key.lower():
                    key = key.rsplit(' ', 1)[0]
                recInfo['Other Resource Type'] = title_case(key)
                # Keep the whole string
                recInfo['Other Resource Amount'] = value   
    # Add dictionary to dataframe
    for key in recInfo:
        df.loc[recID, key] = recInfo[key]
    recID += 1
print("done")

print("Analyzing recommendations...", end ="")
## Calculate on columns
# Calculate payback period
df['Payback Period'] = df['Implementation Cost'] / df['Annual Cost Savings']
# Convert electricity to MMBtu
df['Electricity (MMBtu)'] = df['Electricity (kWh)']* 0.003413/0.33
# Sort df by payback period
df = df.sort_values(by=['Payback Period'])

## Format energy savings strings
for index, row in df.iterrows():
    ST = ""
    SV = ""
    if pd.notna(row['Electricity (kWh)']):
        ST = ST + "Electricity" + '\n\n'
        SV = SV + locale.format_string('%d',row['Electricity (kWh)'], grouping=True) + ' kWh' + '\n'
        SV = SV + '(' + locale.format_string('%d',row['Electricity (MMBtu)'], grouping=True) + ' MMBtu)' + '\n'
    if pd.notna(row['Demand (kW)']):
        ST = ST + "Demand" + '\n'
        SV = SV + locale.format_string('%d',row['Demand (kW)'], grouping=True) + ' kW' + '\n'
    if pd.notna(row['Natural Gas (MMBtu)']):
        ST = ST + "Natural Gas" + '\n'
        SV = SV + locale.format_string('%d',row['Natural Gas (MMBtu)'], grouping=True)  + ' MMBtu' + '\n'
    if pd.notna(row['Other Energy Type']):
        ST = ST + row['Other Energy Type'] + '\n'
        SV = SV + locale.format_string('%d',row['Other Energy Amount'], grouping=True)  + ' MMBtu' '\n'
    if pd.notna(row['Other Resource Type']):
        ST = ST + row['Other Resource Type'] + '\n'
        SV = SV + row['Other Resource Amount'] + '\n'
    ST = ST.rstrip('\n')
    SV = SV.rstrip('\n')
    df.at[index, 'Savings Type'] = ST
    df.at[index, 'Savings Value'] = SV

## Summation statistics
# Filter Recommendationss
recData = df[df['isAdditional'] == False]
# Reorder index
recData = recData.reset_index(drop=True)
# Recommendations statistics
EkWh = recData['Electricity (kWh)'].sum(axis=0, skipna=True)
EMMBtu = recData['Electricity (MMBtu)'].sum(axis=0, skipna=True)
NMMBtu = recData['Natural Gas (MMBtu)'].sum(axis=0, skipna=True)
OMMBtu = recData['Other Energy Amount'].sum(axis=0, skipna=True)
# Add up all energy in MMBtu
iac.MMBtu = round(EMMBtu + NMMBtu + OMMBtu)
# Calculate CO2
if iac.FuelType == "Natural Gas":
    iac.FuelCO2 = 53
    iac.CO2 = round((iac.FuelCO2 * NMMBtu + 0.315 * EkWh)/1000)
elif iac.FuelType == "Propane":
    iac.FuelCO2 = 61.7
    iac.CO2 = round((iac.FuelCO2 * OMMBtu + 0.315 * EkWh)/1000)
elif iac.FuelType == "Fuel Oil #2":
    iac.FuelCO2 = 73.51
    iac.CO2 = round((iac.FuelCO2 * OMMBtu + 0.315 * EkWh)/1000)
# Add up all cost
iac.ACS = recData['Annual Cost Savings'].sum(axis=0, skipna=True)
iac.IC = recData['Implementation Cost'].sum(axis=0, skipna=True)
# Payback period in number
iac.PB = math.ceil(iac.IC / iac.ACS * 10) / 10
# Payback period in formatted string
iac.PBstr = payback(iac.ACS, iac.IC)
print("done")

print("Reformatting recommendations...", end ="")
subtitleList = ["Recommended Actions","Summary of Estimated Savings and Implementation Costs","Current Practice and Observations","Anticipated Savings","Implementation Costs","Implementation Cost References"]
## Reformatting Recommendations
for index, row in recData.iterrows():
    doc = Document(os.path.join('Recommendations', row['File Name']))
    # Change title and make it upper case
    doc.paragraphs[0].text = "Recommendation "+ str(index+1) + ': ' + title_case(row['Description'])
    # Enforce Heading 1
    try:
        doc.paragraphs[0].style = doc.styles['Heading 1']
    except:
        doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        doc.paragraphs[0].style = doc.styles['Heading 1']
    # Enforce subtitle to be Subtitle1
    # This style is already defined in Introduction.docx
    for paragraph in doc.paragraphs:
        for subtitle in subtitleList:
            if paragraph.text == subtitle or paragraph.text == subtitle[:-1]:
                try:
                    paragraph.style = doc.styles['Subtitle1']
                except:
                    doc.styles.add_style('Subtitle1', WD_STYLE_TYPE.PARAGRAPH)
                    paragraph.style = doc.styles['Subtitle1']
    # Save file with sorted filename
    doc.save(os.path.join('Recommendations', 'Sorted', 'Rec'+ str(index+1) + '.docx'))
print("done")

# Check if there's at least 1 additional recommendation
additional = df['isAdditional'].any()
if additional:
    print("Analyzing additional recommendations...", end ="")
    # Filter additional
    addData = df[df['isAdditional'] == True]
    # Reorder index
    addData = addData.reset_index(drop=True)
    # Additional statistics
    EMMBtu = addData['Electricity (MMBtu)'].sum(axis=0, skipna=True)
    NMMBtu = addData['Natural Gas (MMBtu)'].sum(axis=0, skipna=True)
    OMMBtu = addData['Other Energy Amount'].sum(axis=0, skipna=True)
    # Add up all energy
    iac.AddMMBtu = round(EMMBtu + NMMBtu + OMMBtu)
    # Add up all cost
    iac.AddACS = addData['Annual Cost Savings'].sum(axis=0, skipna=True)
    iac.AddIC = addData['Implementation Cost'].sum(axis=0, skipna=True)
    # Payback period in number
    iac.AddPB = round(iac.AddIC / iac.AddACS, 1)
    print("done")

    print("Reformatting additional recommendations...", end ="")
    # Modify the title of the additional recommendation docx
    for index, row in addData.iterrows():
        doc = Document(os.path.join('Recommendations', row['File Name']))
        # Change title and make it upper case
        doc.paragraphs[0].text = "Additional Recommendation"+ str(index+1) + ': ' + title_case(row['Description'])
        # Enforce Heading 1
        try:
            doc.paragraphs[0].style = doc.styles['Heading 1']
        except:
            doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            doc.paragraphs[0].style = doc.styles['Heading 1']
        # Enforce subtitle to be Subtitle1
        # This style is already defined in Introduction.docx
        for paragraph in doc.paragraphs:
            for subtitle in subtitleList:
                if paragraph.text == subtitle:
                    try:
                        paragraph.style = doc.styles['Subtitle1']
                    except:
                        doc.styles.add_style('Subtitle1', WD_STYLE_TYPE.PARAGRAPH)
                        paragraph.style = doc.styles['Subtitle1']
        # Save file with sorted filename
        doc.save(os.path.join('Recommendations', 'Sorted', 'Add'+ str(index+1) + '.docx'))
    print("done")

print("Parsing plant information...", end ="")
## Compiler.json5 Calculations
# Report date = today or 60 days after assessment, which ever is earlier
VD = datetime.datetime.strptime(iac.VDATE, '%B %d, %Y')
RDATE = min(datetime.datetime.today(), VD + datetime.timedelta(days=60))
if platform.system() == 'Windows':
    iac.RDATE = datetime.datetime.strftime(RDATE, '%B %#d, %Y')
else: # macOS or Linux
    iac.RDATE = datetime.datetime.strftime(RDATE, '%B %-d, %Y')

# Sort participant and contributor name list
iac.PARTlist.sort(key=lambda x: x.rsplit(' ', 1)[1])
PART=""
for name in iac.PARTlist:
    PART  = PART + name + '\n'
iac.PART = PART.rstrip('\n')
iac.pop('PARTlist')

iac.CONTlist.sort(key=lambda x: x.rsplit(' ', 1)[1])
CONT=""
for name in iac.CONTlist:
    CONT  = CONT + name + '\n'
iac.CONT = CONT.rstrip('\n')
iac.pop('CONTlist')
print("done")

# products in different cases
iac.PRODTitle = iac.PROD.title()
iac.PRODlower = iac.PROD.lower()

## Format strings
# set electricity cost to 3 digits accuracy
iac = dollar(['EC'],iac,3)
# set the natural gas and demand to 2 digits accuracy
iac = dollar(['DC', 'FC'],iac,2)
# set the rest to integer
varList = ['ACS', 'IC', 'TotalECost', 'TotalFCost', 'TotalCost']
if additional:
    varList.extend(['AddACS', 'AddIC'])
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

## Load introduction template
docIntro = Document(os.path.join('Report', 'Introduction.docx'))

# Add rows to Recommendation table (Should be the 3rd table)
print("Writing recommendation table...", end ="")
locale._override_localeconv={'frac_digits':0}
recTable = docIntro.tables[2]
for index, row in recData.iterrows():
    recRow = recTable.rows[index+1].cells
    # Add ARC No.
    recRow[0].text = 'Rec. ' + str(index+1) + '\n' + row['ARC No.']
    recRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add description
    recRow[1].text = row['Description']
    recRow[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add savings type
    recRow[2].text = row['Savings Type']
    recRow[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add savings value
    recRow[3].text = row['Savings Value']
    recRow[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add annual cost savings
    recRow[4].text = locale.currency(row['Annual Cost Savings'], grouping=True)
    recRow[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Add implementation cost
    recRow[5].text = locale.currency(row['Implementation Cost'], grouping=True)
    recRow[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Add payback period
    pb = row['Payback Period']
    if pb == 0:
        recRow[6].text = "Immediate"
    else:
        recRow[6].text = str(math.ceil(pb * 10) / 10)
    recRow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set 3pt before and after paragraph
    for col in range(0,7):
        recRow[col].paragraphs[0].paragraph_format.space_before = shared.Pt(3)
        recRow[col].paragraphs[0].paragraph_format.space_after = shared.Pt(3)
# Delete unused rows (Currectly row 1-15 are empty)
for index in reversed(range(len(recData), 15)):
    recTable._tbl.remove(recTable.rows[index+1]._tr)
print("done")

if additional:
    # Add rows to additional recommendation table (Should be the 4th table)
    print("Writing Additional Recommendation table...", end ="")
    locale._override_localeconv={'frac_digits':0}
    addTable = docIntro.tables[3]
    for index, row in addData.iterrows():
        addRow = addTable.rows[index+1].cells
        # Add ARC No.
        addRow[0].text = 'Add. Rec. ' + str(index+1) + '\n' + row['ARC No.']
        addRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add description
        addRow[1].text = row['Description']
        addRow[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Add savings type
        addRow[2].text = row['Savings Type']
        addRow[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add savings value
        addRow[3].text = row['Savings Value']
        addRow[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add annual cost savings
        addRow[4].text = locale.currency(row['Annual Cost Savings'], grouping=True)
        addRow[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add implementation cost
        addRow[5].text = locale.currency(row['Implementation Cost'], grouping=True)
        addRow[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add payback period
        pb = row['Payback Period']
        addRow[6].text = str(math.ceil(pb * 10) / 10)
        addRow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set 3pt before and after paragraph
    for col in range(0,7):
        addRow[col].paragraphs[0].paragraph_format.space_before = shared.Pt(3)
        addRow[col].paragraphs[0].paragraph_format.space_after = shared.Pt(3)
    # Delete unused rows (Currectly row 1-5 are empty)
    for index in reversed(range(len(addData), 5)):
        addTable._tbl.remove(addTable.rows[index+1]._tr)
    print("done")
else:
    # delete this table
    docIntro._body._body.remove(docIntro.tables[3]._tbl)

# Remove Add blocks if no Additional
docx_blocks(docIntro, ADD = additional)

# Replacing keys
print("Replacing keys in introduction...", end ="")
docx_replace(docIntro, **iac)
print("done")

# Save introduction
filename_intro = iac.LE + '-intro.docx'
docIntro.save(filename_intro)

## Load backgroud template
docBackground = Document(os.path.join('Report', 'Background.docx'))

# Replacing keys
print("Replacing keys in background...", end ="")
docx_replace(docBackground, **iac)
print("done")

# Save background
filename_back = iac.LE + '-back.docx'
docBackground.save(filename_back)

## Load energy bill analysis template
docEnergy = Document(os.path.join('Report', 'Energy.docx'))

# Add energy chart images
print("Adding energy chart images...", end ="")
# If on macOS
if chartPath == os.path.join('Energy Charts', 'Energy Charts.fld'):
    add_image(docEnergy, '#EUChart', os.path.join(chartPath, "image001.png"), shared.Inches(6))
    add_image(docEnergy, '#ECChart', os.path.join(chartPath, "image004.png"), shared.Inches(6))
    add_image(docEnergy, '#DUChart', os.path.join(chartPath, "image002.png"), shared.Inches(6))
    add_image(docEnergy, '#DCChart', os.path.join(chartPath, "image005.png"), shared.Inches(6))
    add_image(docEnergy, '#FUChart', os.path.join(chartPath, "image003.png"), shared.Inches(6))
    add_image(docEnergy, '#FCChart', os.path.join(chartPath, "image006.png"), shared.Inches(6))
    add_image(docEnergy, '#PieUChart', os.path.join(chartPath, "image007.png"), shared.Inches(6))
    add_image(docEnergy, '#PieCChart', os.path.join(chartPath, "image008.png"), shared.Inches(6))
    add_image(docEnergy, '#TotalChart', os.path.join(chartPath, "image009.png"), shared.Inches(9))
# If on Windows
elif chartPath == os.path.join('Energy Charts', 'Energy Charts_files'):
    add_image(docEnergy, '#EUChart', os.path.join(chartPath, "image001.png"), shared.Inches(6))
    add_image(docEnergy, '#ECChart', os.path.join(chartPath, "image002.png"), shared.Inches(6))
    add_image(docEnergy, '#DUChart', os.path.join(chartPath, "image003.png"), shared.Inches(6))
    add_image(docEnergy, '#DCChart', os.path.join(chartPath, "image005.png"), shared.Inches(6))
    add_image(docEnergy, '#FUChart', os.path.join(chartPath, "image006.png"), shared.Inches(6))
    add_image(docEnergy, '#FCChart', os.path.join(chartPath, "image007.png"), shared.Inches(6))
    add_image(docEnergy, '#PieUChart', os.path.join(chartPath, "image009.png"), shared.Inches(6))
    add_image(docEnergy, '#PieCChart', os.path.join(chartPath, "image011.png"), shared.Inches(6))
    add_image(docEnergy, '#TotalChart', os.path.join(chartPath, "image013.png"), shared.Inches(9))
print("done")

# Fill in energy chart tables from Energy Charts.xlsx
print("Adding energy chart tables...", end ="")
# Read electricity table from B6 to I19
edf = pd.read_excel(os.path.join('Energy Charts', 'Energy Charts.xlsx'), sheet_name="Raw Data", skiprows = 5, nrows=13, usecols = 'B:I')
# Read fuel table from K6 to N19
fdf = pd.read_excel(os.path.join('Energy Charts', 'Energy Charts.xlsx'), sheet_name="Raw Data", skiprows = 5, nrows=13, usecols = 'K:N')

# Add rows to electricity table (Should be the 1st table)
eTable = docEnergy.tables[0]
for index, row in edf.iterrows():
    eRow = eTable.rows[index+3].cells
    # Add Month
    eRow[0].text = edf.iloc[(index, 0)]
    eRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col in range(1,8):
        # Add interger with thousand separator
        eRow[col].text = locale.format_string('%d',round(edf.iloc[(index, col)]), grouping=True)
        eRow[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Bold the last row
    if index == 12:
        for col in range(0,8):
            eRow[col].paragraphs[0].runs[0].bold = True

# Add rows to fuel table (Should be the 2nd table)
fTable = docEnergy.tables[1]
for index, row in fdf.iterrows():
    fRow = fTable.rows[index+3].cells
    # Add Month
    fRow[0].text = fdf.iloc[(index, 0)]
    fRow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col in range(1,4):
        # Add interger with thousand separator
        fRow[col].text = locale.format_string('%d',round(fdf.iloc[(index, col)]), grouping=True)
        fRow[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Bold the last row
    if index == 12:
        for col in range(0,4):
            fRow[col].paragraphs[0].runs[0].bold = True
print("done")
# Replacing keys
print("Replacing keys in energy charts...", end ="")
docx_replace(docEnergy, **iac)
print("done")
# Save energy charts
filename_energy = iac.LE + '-energy.docx'
docEnergy.save(filename_energy)

print("Combining all docs...", end ="")
# List of docs to combine
docList = [os.path.join('Report', 'ToC.docx')]
for RecLength in range(1, len(recData)+1):
    docList.append(os.path.join('Recommendations', 'Sorted','Rec' + str(RecLength) + '.docx'))
if additional:
    docList.append(os.path.join('Report', 'Add.docx'))
    for AddLength in range(1, len(addData)+1):
        docList.append(os.path.join('Recommendations', 'Sorted','Add' + str(AddLength) + '.docx'))
else:
    pass

# Combine all docx files
main = Document(filename_intro)
main.add_page_break()
composer = Composer(main)
for i in range(0, len(docList)):
    doc_add = Document(docList[i])
    doc_add.add_page_break()
    composer.append(doc_add)
composer.append(Document(filename_back))
composer.append(Document(filename_energy))
filename = iac.LE +'.docx'
composer.save(filename)

# delete temp files
os.remove(filename_intro)
os.remove(filename_back)
os.remove(filename_energy)
print("done")

# Open the combined docx file
doc = Document(filename)
# Change the orientation of the last section to landscape
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

# Save final report
doc.save(filename)
print(filename + " is finished.")

# Caveats
caveat("Please fix title case in ToC (AR/ARR and conjunctions).")
caveat("Please select all (Ctrl+A) then refresh TWICE (F9) ToC, list of tables/figures.")
caveat("Please select list of tables/figures then set to NO BOLD.")
caveat("Please manually add Process Description, Major Equipment, Current Best Practices, and plant layout image.")
