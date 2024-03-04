"""
This script is used to generate the IAC recommendation for Install Motion Sensor
"""

import json5, sys, os, num2words
from docx import Document
from easydict import EasyDict
from python_docx_replace import docx_replace, docx_blocks
sys.path.append(os.path.join('..', '..'))
from Shared.IAC import *
from docxcompose.composer import Composer
import numpy as np

# Load utility cost
jsonDict = json5.load(open(os.path.join('..', '..', 'Utility.json5')))
# Load database
jsonDict.update(json5.load(open('database.json5')))
# Convert to easydict
iac = EasyDict(jsonDict)

## Validate the length of all lists
N = iac.N
for i in iac:
    if isinstance(iac[i], list):
        # If the length is not N, throw an error
        if len(iac[i]) != N:
            raise Exception('Length of {0} is not {1}.'.format(i, N))
## Covert to numpy array for element-wise operations
nplist = ['LED', 'CFW', 'HR', 'DY', 'WK', 'FR']
for i in nplist:
    iac[i] = np.array(iac[i])

## Constants
# Conversion constant; W/kW
C1 = 1000 

## Calculations
# Operating Hours
iac.OH = iac.HR * iac.DY * iac.WK

## Savings
# Annual electricity savings
iac.ESi = np.rint((iac.LED * iac.CFW * iac.OH * (100/100 - iac.FR/100))/ C1)
# Total Energy Savings
iac.ES = np.sum(iac.ESi)
# Annual cost savings
iac.ACS = iac.ES * iac.EC

## Implementation cost Estimate
# Total cost for all sensors
iac.TCOST = iac.COST * N
iac.TLABOR = iac.LABOR * N
iac.IC = iac.TCOST + iac.TLABOR
# Rebate
iac.PB = payback(iac.ACS.item(), iac.IC)

# TItle Converter
iac.TLOC = [''] * N
for i in range(N):
  iac.TLOC[i] = iac.LOC[i].title()

# Number to words
iac.NUM = num2words.num2words(N)

## Format strings
# set to 3 digits accuracy
iac = dollar(['EC'],iac,3)
# set the rest to integer
varList = ['COST', 'LABOR', 'TCOST', 'TLABOR', 'ACS', 'IC']
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

# Create document for each area
for i in range(N):
  iacsub = EasyDict()
  iacsub.i = str(i+1)
  # For any list or ndarray in iac, add corresponding values to iacsub
  for j in iac:
    if isinstance(iac[j], list) or isinstance(iac[j], np.ndarray):
      iacsub[j] = iac[j][i]
  # Import individual area template
  doc = Document('template 2.docx')
  # Replacing keys
  docx_replace(doc, **iacsub)
  # Save file as temp{i}.docx
  doc.save('tmp'+iacsub.i+'.docx')

# Import opening template
doc = Document('template 1.docx')
# Replacing keys
docx_replace(doc, **iac)
# Save file as temp0.docx
doc.save('tmp0.docx')

# Assemble ESSum
iac.ESSum = iac.ESi[0] + ' kWh/yr'
for i in range(1, N):
   iac.ESSum += ' + ' + iac.ESi[i] + ' kWh/yr'

# Import ending template
doc = Document('template 3.docx')
# Replacing keys
docx_replace(doc, **iac)
# Save file as temp{N+1}.docx
doc.save('tmp'+str(N+1)+'.docx')

# Combine all docx files
master = Document("tmp0.docx")
composer = Composer(master)
for i in range(N+1):
    doc_tmp = Document('tmp'+str(i+1)+'.docx')
    composer.append(doc_tmp)

savefile(composer, str(iac.REC))

# delete temp files
for i in range(N+2):
    filename = 'tmp'+str(i)+'.docx'
    os.remove(filename)
# Caveats
caveat("Please change implementation cost references if necessary.")