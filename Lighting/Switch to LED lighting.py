"""
This script is used to generate the IAC recommendation for Switch to LED lighting.
"""

import json5, sys, os
from docx import Document
from easydict import EasyDict
from python_docx_replace import docx_replace, docx_blocks
sys.path.append('..') 
from Shared.IAC import *
from docxcompose.composer import Composer
import numpy as np

# Load config file and convert everything to EasyDict
jsonDict = json5.load(open('Switch to LED lighting.json5'))
jsonDict.update(json5.load(open(os.path.join('..', 'Utility.json5'))))
iac = EasyDict(jsonDict)

# Validate the length of all lists
N = iac.N
for i in iac:
    if isinstance(iac[i], list):
        # If the length is not N, throw an error
        if len(iac[i]) != N:
            raise Exception('Length of {0} is not {1}.'.format(i, N))

# Calculations
# Covert to numpy array for element-wise operations
nplist = ['CN', 'CPR', 'CHR', 'CDY', 'CWK', 'PN', 'PPR', 'PHR', 'PDY', 'PWK', 'BP', 'BL', 'CF']
for i in nplist:
    iac[i] = np.array(iac[i])

# Calculate operating hours
iac.COH = iac.CHR * iac.CDY * iac.CWK
iac.POH = iac.PHR * iac.PDY * iac.PWK
# Calculate electricity savings
iac.ESi = np.rint((iac.CN * iac.CPR * iac.COH - iac.PN * iac.PPR * iac.POH) / 1000.0).astype(np.int64)
iac.ES = np.sum(iac.ESi)
iac.ECS = np.rint(iac.ES * iac.EC).astype(np.int64)
# Calculate demand savings
iac.DSi = np.rint((iac.CN * iac.CPR - iac.PN * iac.PPR) * (iac.CF/100) * 12.0 / 1000.0).astype(np.int64)
iac.DS = np.sum(iac.DSi)
iac.DCS = np.rint(iac.DS * iac.DC).astype(np.int64)

# Calculate bulb cost
iac.BCi = np.rint(iac.PN * iac.BP).astype(np.int64)
iac.BC = np.sum(iac.BCi)
# Calculate labor cost
iac.LCi = np.rint(iac.CN * iac.BL).astype(np.int64)
iac.LC = np.sum(iac.LCi)
# Calculate implementation cost
iac.LN = np.sum(iac.CN)
iac.MSC = iac.MSN * iac.MSPL
iac.IC = iac.MSC + iac.BC + iac.LC
iac.ACS = iac.ECS + iac.DCS

# Rebate
iac.RB = round(iac.ES * iac.RR)
iac.MRB = min(iac.RB, iac.IC/2)
iac.MIC = iac.IC - iac.MRB
iac.PB = payback(iac.ACS.item(), iac.MIC.item())

# Combine words
iac.AREAS = combine_words(iac.AREA)
# Take an example of the previous area
iac.PREV1 = iac.PREV[0]

# Motion sensor
if iac.MSN == 0:
    MS = False
else:
    MS = True

## Format strings
# set electricity cost / rebate to 3 digits accuracy
iac = dollar(['EC', 'RR'],iac,3)
# set the natural gas and demand to 2 digits accuracy
iac = dollar(['NGC', 'DC'],iac,2)
# set the rest to integer
varList = ['LR', 'MSPL', 'ECS', 'DCS', 'ACS', 'MSC', 'BC', 'LC', 'IC', 'RB', 'MRB', 'MIC']
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

# Create document for each area
for i in range(N):
    iacsub = EasyDict()
    iacsub.i = str(i+1)
    # For any list or ndarray in iac, add corresponding values to iacsub
    for j in iac:
       if isinstance(iac[j], list) or isinstance(iac[j], np.ndarray):
           iacsub[j] = iac[j][i]

    # Import individual area template
    doc = Document('Switch to LED lighting 2.docx')

    # Replacing keys
    docx_replace(doc, **iacsub)
    # Save file as temp{i}.docx
    doc.save('temp'+iacsub.i+'.docx')

# Import opening template
doc = Document('Switch to LED lighting 1.docx')
# Replacing keys
docx_replace(doc, **iac)
# Save file as temp0.docx
doc.save('temp0.docx')

# Assemble ESSum and ESSum
iac.ESSum = iac.ESi[0] + ' kWh/yr'
iac.DSSum = iac.DSi[0] + ' kW/yr'
for i in range(1,N):
    iac.ESSum += ' + ' + iac.ESi[i] + ' kWh/yr'
    iac.DSSum += ' + ' + iac.DSi[i] + ' kW/yr'

# Import ending template
doc = Document('Switch to LED lighting 3.docx')
# Motion sensors block
docx_blocks(doc, ms = MS)
# Multi areas block
if N == 1:
    docx_blocks(doc, single = True)
    docx_blocks(doc, multi = False)
else:
    docx_blocks(doc, single = False)
    docx_blocks(doc, multi = True)
iac.INSTALL = []
for i in range(N):
    tmpstr = f"a {iac.PPR[i]} W linear LED bulb "
    tmpstr += f"costs about {iac.BP[i]} plus "
    tmpstr += f"{iac.BL[i]} labor to install"
    # captialize the first letter of the first sentence
    if i==0:
        tmpstr = tmpstr[0].capitalize() + tmpstr[1:]
    iac.INSTALL.append(tmpstr)
iac.INSTALL = combine_words(iac.INSTALL)
# Replacing keys
docx_replace(doc, **iac)
# Save file as temp{N+1}.docx
doc.save('temp'+str(N+1)+'.docx')

# Combine all docx files
master = Document("temp0.docx")
composer = Composer(master)
for i in range(N+1):
    doc_temp = Document('temp'+str(i+1)+'.docx')
    composer.append(doc_temp)
filename = "AR" + iac.AR + ".docx"
composer.save(os.path.join("..", "ARs", filename))
# delete temp files
for i in range(N+2):
    filename = 'temp'+str(i)+'.docx'
    os.remove(filename)

# Caveats
caveat("Please change implementation cost references if necessary.")