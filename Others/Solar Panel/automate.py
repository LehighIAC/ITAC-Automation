"""
This script is used to generate the IAC recommendation for Install Solar Panels.
"""

import json5, sys, os, locale
from docx import Document
from easydict import EasyDict
from docx.enum.text import WD_ALIGN_PARAGRAPH
from python_docx_replace import docx_replace
sys.path.append(os.path.join('..', '..')) 
from Shared.IAC import *
import requests, datetime

# Load utility cost
jsonDict = json5.load(open(os.path.join('..', '..', 'Utility.json5')))
# Load database
jsonDict.update(json5.load(open('database.json5')))
# Convert to easydict
iac = EasyDict(jsonDict)

# Different template for PA/NJ
if iac.ST == "PA":
    template = "template - PA.docx"
    iac.AMV = iac.AMVPA
elif iac.ST == "NJ":
    template = "template - NJ.docx"
    iac.AMV = iac.AMVNJ
else:
    pass

# Calculations
# Avaialble space ft2
iac.AS = round(iac.RS * iac.ASR / 100)
# Capacity kW
iac.CAP = round(iac.AS / 100)
# Approx. energy savings, kWh
iac.AES = iac.CAP * 1200

# PVWatts API
parameters = {
'format': 'json',
'api_key': iac.api,
'system_capacity': iac.CAP,
'module_type': 0,
'losses': 14.08,
'array_type': 0,
'tilt': 20,
'azimuth': 180,
'address': iac.ZIP,
}

solard_monthly = [2.49, 3.42, 4.24, 5.07, 5.73, 5.89, 6.30, 5.60, 4.72, 3.64, 2.96, 1.98]
ac_monthly = [10762, 13137, 17398, 19346, 21770, 21224,23189, 20814, 17107, 14396, 11782,8576]
try:
  response = requests.request('GET', 'https://developer.nrel.gov/api/pvwatts/v8.json', params=parameters)
  response.raise_for_status()
  PVresults = response.json()
  iac.ES = round(PVresults.get('outputs').get('ac_annual'))
  # read solard_monthly and ac_monthly
  solard_monthly = PVresults.get('outputs').get('solrad_monthly')
  ac_monthly = PVresults.get('outputs').get('ac_monthly')
except:
  print(response.status_code)
  print('PVWatts API error. Please look up the annual energy savings manually on PVWatts website')
  # input number
  iac.ES = int(input('Manually input annual energy savings (kWh): '))

iac.ACSel = round(iac.ES * iac.EC)
iac.credits = round(iac.ES / 1000)
iac.ACSsr = round(iac.AMV * iac.credits)
iac.ACS = iac.ACSel + iac.ACSsr

# Implementation cost
iac.IC = round(iac.CAP * iac.PPW * 1000)
iac.ITC = round(iac.IC * iac.ITCR / 100)
iac.MIC = iac.IC - iac.ITC
iac.PB = payback(iac.ACS, iac.MIC)
iac.CM = datetime.datetime.now().strftime('%B %Y')

## Format strings
# set electricity cost / rebate to 3 digits accuracy
iac = dollar(['EC'],iac,3)
# set the natural gas and demand to 2 digits accuracy
iac = dollar(['NGC', 'DC', 'PPW'],iac,2)
# set the rest to integer
varList = ['LR', 'MIC', 'IC', 'ITC', 'AMV', 'ACSel', 'ACSsr', 'ACS']
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

doc = Document(template)

# Replacing keys
docx_replace(doc, **iac)

# Set local to US
locale.setlocale(locale.LC_ALL, 'en_US')
# Fill in the second table
table = doc.tables[1]
for i in range(12):
    table.cell(i+1, 1).text = str(round(solard_monthly[i],2))
    table.cell(i+1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(i+1, 2).text = locale.format_string('%d',round(ac_monthly[i]), grouping=True)
    table.cell(i+1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
table.cell(13, 1).text = str(round(sum(solard_monthly)/12,2))
table.cell(13, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
table.cell(13, 1).paragraphs[0].runs[0].bold = True
table.cell(13, 2).text = locale.format_string('%d',round(sum(ac_monthly)), grouping=True)
table.cell(13, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
table.cell(13, 2).paragraphs[0].runs[0].bold = True

savefile(doc, iac.REC, add=True)

# Caveats
caveat("Please check if the grabbed info is correct.")